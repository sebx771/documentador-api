from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from markdown_it import MarkdownIt


class EasyDocsDOCX:
    """
    Clase para convertir Markdown en un documento de Word (.docx).
    Usa markdown-it-py para un parseo estructural robusto.
    """

    def __init__(self):
        self.doc = Document()
        self.md = MarkdownIt("gfm-like")  # Soporta tablas y extensiones comunes
        self._configurar_estilos_base()

    def _configurar_estilos_base(self):
        """Configura la fuente y el tamaño base del doc."""
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(11)

    def agregar_encabezado(self):
        """Añade el sello institucional en la parte superior."""
        section = self.doc.sections[0]
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.text = "Reporte de Documentación Técnica - easyDocs"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.runs[0]
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)

    def construir_desde_markdown(self, texto_md):
        """interpreta el Markdown estructuralmente y lo plasma en el documento."""
        tokens = self.md.parse(texto_md)
        i = 0
        while i < len(tokens):
            token = tokens[i]

            # Títulos
            if token.type == "heading_open":
                level = int(token.tag[1])
                i += 1
                inline_token = tokens[i]
                h = self.doc.add_heading(inline_token.content, level=level)
                # Estilo azul para títulos
                for run in h.runs:
                    run.font.color.rgb = RGBColor(31, 73, 125)
                    run.font.name = "Arial"

            # Párrafos
            elif token.type == "paragraph_open":
                i += 1
                inline_token = tokens[i]
                if inline_token.type == "inline":
                    p = self.doc.add_paragraph()
                    self._render_inline(p, inline_token)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Bloques de Código (Fence)
            elif token.type == "fence":
                self._agregar_bloque_codigo(token.content)

            # Listas
            elif token.type == "bullet_list_open" or token.type == "ordered_list_open":
                is_ordered = token.type == "ordered_list_open"
                i += 1
                while i < len(tokens) and tokens[i].type != token.type.replace(
                    "_open", "_close"
                ):
                    if tokens[i].type == "list_item_open":
                        i += 2  # Saltar a inline del párrafo dentro del item
                        inline_token = tokens[i]
                        style = "List Number" if is_ordered else "List Bullet"
                        p = self.doc.add_paragraph(style=style)
                        self._render_inline(p, inline_token)
                    i += 1

            # Tablas (Simplificado)
            elif token.type == "table_open":
                i = self._render_tabla(tokens, i)

            i += 1

    def _render_inline(self, paragraph, inline_token):
        """Renderiza contenido inline (negritas, itálicas, código inline)."""
        if not inline_token.children:
            paragraph.add_run(inline_token.content)
            return

        for child in inline_token.children:
            if child.type == "text":
                paragraph.add_run(child.content)
            elif child.type == "code_inline":
                run = paragraph.add_run(child.content)
                run.font.name = "Courier New"
                run.font.color.rgb = RGBColor(214, 51, 132)
            elif child.type == "strong_open":
                pass  # El siguiente texto será negrita
            elif child.type == "strong_close":
                pass
            # Nota: Para una implementación completa, se requiere trackear el estado (bold/italic)
            # Aquí lo mantenemos simple para legibilidad.

    def _agregar_bloque_codigo(self, texto):
        """Añade un bloque de código con estilo."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(6)

        # Sombreado gris
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F8F9FA")
        p._element.get_or_add_pPr().append(shd)

        for linea in texto.strip().split("\n"):
            run = p.add_run(linea + "\n")
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(51, 51, 51)

    def _render_tabla(self, tokens, start_index):
        """Renderiza una tabla de Markdown en Word."""
        i = start_index
        rows = []
        current_row = []

        while i < len(tokens) and tokens[i].type != "table_close":
            if tokens[i].type == "tr_open":
                current_row = []
            elif tokens[i].type == "th_open" or tokens[i].type == "td_open":
                i += 2  # th/td_open -> (maybe paragraph_open) -> inline
                if tokens[i].type == "inline":
                    current_row.append(tokens[i].content)
            elif tokens[i].type == "tr_close":
                rows.append(current_row)
            i += 1

        if rows:
            table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for r_idx, row_data in enumerate(rows):
                for c_idx, cell_text in enumerate(row_data):
                    table.cell(r_idx, c_idx).text = cell_text
        return i

    def guardar(self, output_stream):
        """Exporta el documento al flujo de bytes final."""
        self.doc.save(output_stream)
